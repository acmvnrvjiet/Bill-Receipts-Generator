import subprocess
import sys
import os
def install(package):   #defining function "install"
    subprocess.check_call([sys.executable,"-m","pip","install",package])
install('openpyxl')   #function calling
install('docx')
install('docx2pdf')
install('docx.shared')
install('shutil')
from docx.shared import Pt
#from docx.shared import Inches
from docx2pdf import convert
#docx2pdf for converting document to pdf 
from docx import Document
#docx for handling docx files
from shutil import copy2
# shutil for file operatins, copy2 for copying the file from 
from openpyxl import load_workbook
# openpyxl package for accessing excel sheets
from tkinter.filedialog import askopenfile


try:
    file = askopenfile(title='Select the Workbook', mode='r', filetypes=[
        ('Microsoft Excel', '.xlsx .xlsx .xlsm .xltx .xltm')])
#this displays a dialog box that allows us to select the file(excel workbook)
except:
    {}
#try except block incase any exceptions or choosing wrong type of file

if file is not None:   #if file is not empty(i.e it has some data)
    dirpath = os.path.dirname(file.name)  
    wbname = os.path.basename(file.name)
else:
    sys.exit()

filepath = file.name
# filepath variable contains the path of the members info excel sheet
# r'' makes d string in raw format so dat special characters lyk \n,\t will nt b taken into consideration

dirpath = dirpath+'/'+'AckReceipts'
# dirpath variable contains the address of location where we are going to save the Receipts
# creating a folder named AckReceipts

allIdsPath = dirpath+'/All_Receipts'
#creating another folder named All_Receipts 


try:
    os.mkdir(dirpath)
    os.mkdir(allIdsPath)
except:
    {}
# os.mkdir(filepath) is used to create a folder/dir where we are saving the receipts
# its written in try-except block so that if the folder is already present then exception is generated
# but will not cause any problem to the execution of the program   

wb = load_workbook(filepath)
# wb variable stores the workbook of the opened excel file
# load_workbook opens the workbook

print(end='\n\n')
print(wbname, end='\n\n')

l = [1, 2, 11, 3, 8, 6, 7, 10, 9, 12]
#this is the list of the column number as in the excel sheets in the order of the cells in the table in the template 

for ws in wb:     # for every worksheet(ws) in the opened workbook(wb):

    print(ws.title)

    wspath = dirpath+'/'+ws.title.strip()    # .strip() is used to remove extra spaces at the end of a string
    # wspath stores the path of another inside folder with the name of the each worksheet 

    try:
        os.mkdir(wspath)    #  it creates another folder/dir inside 'AckReceipts' folder with the name of the worksheet ws
    except:
        {}

    # for every row in the worksheet which starts from 3 to 60
    # as the info of members starts from row no 3
    # at max no. of students could be 60
    for r in range(3, 60):
        cell = ws.cell(row=r, column=1)
        # cell variable selects the cell which is placed at location row=r n column=1

        if cell.value is None:
            break
        # if the first cell(sl_no cell) is empty(i.e row is empty/reached end) the loop will break
        # cell.value gives the data inside the cell

        # rno,fname,lname are variables which store the respective data of a member of row=r    
        rno = str(ws.cell(row=r, column=8).value).strip().upper()  #upper() to change rno to uppercase since it contains alphabets
        fname = ws.cell(row=r, column=4).value.strip().title()     #title() changes first letter in the word to upper case and rest all letters to lower case 
        lname = ws.cell(row=r, column=5).value
        if(lname is None):
            lname = ''        #if last name is empty then store it as empty string 
        else:
            lname = lname.strip().title()

        ws.cell(row=r, column=10).value = ws.cell(
            row=r, column=10).value.strip().lower()  # changing all the letters in the email id to lowercase and storing it back in the same cell
        ws.cell(row=r, column=8).value = rno         
        ws.cell(row=r, column=4).value = fname
        ws.cell(row=r, column=5).value = lname
        # storing the modified rno,fname,lname in the respective cells of row r

        #generating Id nos
        #Use a pattern to create the required ID and write the code down here
        acm_id = 'ID'+rno[8:]

        ws.cell(row=r, column=2).value = acm_id   #storing the ID no that has been created in the cell of column 2 & row r
        ws.cell(row=r, column=14).value = acm_id+'_receipt.pdf'   #storing the receipt pdf name that will be created in the cell of column 14 & row r
        ws.cell(row=r, column=3).value = fname+' '+lname   ##storing the full name in the cell of column 3 & row r

        eachmemberpath = wspath+'/'+acm_id
        # eachmemberpath variable stores the path address where we are going to store the respective person's receipt 
        # it creates a folder with name as acm_id(i.e ID no) inside the sheet name(ex: sheet1) folder

        try:
            os.mkdir(eachmemberpath)
        except:
            {}

        document = Document('Template.docx') #add full path if any error occurs
        #accessing the receipt template(which is in doc format)

        i = 0
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text == '':
                        if(l[i] == 11):
                            x = str(ws.cell(row=r, column=l[i]).value).strip().split()[   # column 11 is date, removing spaces at the end using strip() & 
                                0].split('-')                                             # splitting using split()
                            print(x)
                            date = x[2]+'/'+x[1]+'/'+x[0] 
                            font = cell.paragraphs[0].add_run(date).font
                        else:
                            font = cell.paragraphs[0].add_run(
                                str(ws.cell(row=r, column=l[i]).value).strip()).font
                        font.name = 'Calibri'
                        font.size = Pt(12)
                        i += 1                                         #filling each cell in the table 
        idpath = eachmemberpath+'/'+acm_id+'_receipt'
        document.save(idpath+'.docx')          #saving the document after filling the table

        convert(idpath+'.docx', idpath+'.pdf')    #converting the document to pdf format 
        copy2(idpath+'.pdf', allIdsPath+'/'+acm_id+'_receipt.pdf')     #copying the pdf to the folder "All_Receipts"

        print(acm_id, ' Done.')       #printing in the logs after the job is done
    print()

wb.save(file.name)
print(wbname, 'WorkBook Done.', end='\n\n')
# after the job is done ... a Done statement is printed in the logs...!!!
